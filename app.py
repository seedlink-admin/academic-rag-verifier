"""
📚 전천후 도서/논문 검증 및 연구 조교 (Academic RAG Verifier)
- Google Gemini API + RAG(검색 증강 생성) 기반
- TXT / PDF / EPUB / DOCX 다중 업로드 지원
- PyMuPDF 기반 학술 논문 정밀 파싱, ChromaDB 벡터 저장소
- 환각(Hallucination) 차단 검증 전용 프롬프트
"""

import os
import re
import hashlib
import tempfile
import traceback

import streamlit as st
import fitz  # PyMuPDF
import chromadb
import google.generativeai as genai
from langchain_text_splitters import RecursiveCharacterTextSplitter
from dotenv import load_dotenv

load_dotenv()

# ──────────────────────────────────────────────────────────────
# 기본 설정
# ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="전천후 도서/논문 검증 연구 조교",
    page_icon="📚",
    layout="wide",
)

EMBEDDING_MODEL = "models/text-embedding-004"
HF_FALLBACK_MODELS = ["jhgan/ko-sroberta-multitask", "sentence-transformers/all-MiniLM-L6-v2"]
COLLECTION_NAME = "academic_rag_collection"

# Chunk: 800~1000 토큰 ≈ 3200~4000자(영문) / 한국어 혼용 고려 3400자, Overlap 150토큰 ≈ 600자
CHUNK_SIZE = 3400
CHUNK_OVERLAP = 600
TOP_K = 8

SYSTEM_PROMPT = """당신은 업로드된 서적 및 학술 논문의 팩트만을 검증하는 '엄격한 학술 검증 AI'입니다.
[규칙 1] 반드시 제공된 [참고 문서] 텍스트 안에 명시된 사실만을 기반으로 답변하십시오.
[규칙 2] 만약 사용자가 질문한 내용(예: 출판사, 발행일, 실험 데이터, 인용구 등)이 [참고 문서]에 없다면, 절대 외부 지식이나 추측을 동원하지 말고 정확히 "제공된 문서에서 해당 정보를 찾을 수 없습니다."라고만 답변하십시오.
[규칙 3] 여러 문서가 업로드된 경우, 문서 간의 내용이 일치하는지 또는 상충하는지 교차 검증(Cross-check)하여 답변에 반영하십시오.
[규칙 4] 답변 끝에는 항상 [출처: 문서명, p.페이지 번호 또는 섹션]을 구체적으로 적어주십시오."""

QUICK_ACTIONS = {
    "bib": (
        "📌 서지 정보 추출",
        "업로드된 각 문서의 서지 정보(제목, 저자, 출판사 또는 게재 학술지, 발행일, DOI 등)를 "
        "문서별로 표 형태로 정리해 주세요. 문서에 명시되지 않은 항목은 '문서 내 정보 없음'으로 표기하세요.",
    ),
    "summary": (
        "📊 핵심 요약",
        "업로드된 문서(논문)의 (1) 연구 목적, (2) 연구 방법, (3) 핵심 결론을 각각 한 단락씩, "
        "총 3단락으로 요약해 주세요. 반드시 문서에 명시된 내용만 사용하세요.",
    ),
    "crosscheck": (
        "🔍 논리 교차 검증",
        "현재 업로드된 여러 문서 간의 핵심 주장, 실험 데이터, 결론을 비교하여 서로 일치하는 부분과 "
        "모순되거나 상충하는 부분을 교차 검증(Cross-check)해 주세요. 문서가 1개뿐이라면 해당 문서 "
        "내부의 논리적 일관성(주장-근거-결론)을 검증해 주세요.",
    ),
}


# ──────────────────────────────────────────────────────────────
# 세션 상태 초기화
# ──────────────────────────────────────────────────────────────
def init_session():
    defaults = {
        "messages": [],          # [{"role", "content", "sources"}]
        "docs": {},              # {doc_hash: {"name", "status", "chunks", "pages"}}
        "chroma_client": None,
        "collection": None,
        "api_key_ok": False,
        "pending_prompt": None,
        "embed_backend": None,   # "google" | "hf" — 최초 임베딩 성공 시 확정(차원 혼용 방지)
        "embed_model_name": None,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


init_session()


def get_collection():
    if st.session_state.chroma_client is None:
        st.session_state.chroma_client = chromadb.EphemeralClient()
    if st.session_state.collection is None:
        st.session_state.collection = st.session_state.chroma_client.get_or_create_collection(
            name=COLLECTION_NAME, metadata={"hnsw:space": "cosine"}
        )
    return st.session_state.collection


# ──────────────────────────────────────────────────────────────
# 문서 파싱 (형식별)
# ──────────────────────────────────────────────────────────────
def parse_pdf(file_bytes: bytes):
    """PyMuPDF로 페이지별 텍스트를 손실 없이 추출. 2단 편집 대응을 위해 blocks 정렬 사용."""
    units = []  # [(location_label, text)]
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for i, page in enumerate(doc):
            # sort=True: 다단 편집(2단 컬럼)에서도 읽기 순서를 최대한 보존
            text = page.get_text("text", sort=True)
            text = re.sub(r"[ \t]+", " ", text).strip()
            if text:
                units.append((f"p.{i + 1}", text))
    return units


def parse_txt(file_bytes: bytes):
    for enc in ("utf-8", "cp949", "euc-kr", "latin-1"):
        try:
            text = file_bytes.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = file_bytes.decode("utf-8", errors="ignore")
    return [("전체 텍스트", text.strip())] if text.strip() else []


def parse_docx(file_bytes: bytes):
    import docx as _docx

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(file_bytes)
        path = tmp.name
    try:
        d = _docx.Document(path)
        paras = [p.text for p in d.paragraphs if p.text.strip()]
        # 표 내용도 추출
        for table in d.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paras.append(" | ".join(cells))
        text = "\n".join(paras)
        return [("본문", text)] if text.strip() else []
    finally:
        os.unlink(path)


def parse_epub(file_bytes: bytes):
    from ebooklib import epub, ITEM_DOCUMENT
    from bs4 import BeautifulSoup

    with tempfile.NamedTemporaryFile(suffix=".epub", delete=False) as tmp:
        tmp.write(file_bytes)
        path = tmp.name
    try:
        book = epub.read_epub(path)
        units = []
        idx = 0
        for item in book.get_items():
            if item.get_type() == ITEM_DOCUMENT:
                soup = BeautifulSoup(item.get_content(), "html.parser")
                text = soup.get_text(separator="\n").strip()
                text = re.sub(r"\n{3,}", "\n\n", text)
                if text:
                    idx += 1
                    title_tag = soup.find(["h1", "h2", "h3", "title"])
                    label = (
                        f"섹션 {idx}: {title_tag.get_text().strip()[:40]}"
                        if title_tag and title_tag.get_text().strip()
                        else f"섹션 {idx}"
                    )
                    units.append((label, text))
        return units
    finally:
        os.unlink(path)


def parse_file(name: str, file_bytes: bytes):
    ext = name.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        return parse_pdf(file_bytes)
    if ext == "txt":
        return parse_txt(file_bytes)
    if ext == "docx":
        return parse_docx(file_bytes)
    if ext == "epub":
        return parse_epub(file_bytes)
    raise ValueError(f"지원하지 않는 파일 형식: {ext}")


# ──────────────────────────────────────────────────────────────
# 청킹 + 임베딩 + 벡터 저장
# ──────────────────────────────────────────────────────────────
def chunk_units(doc_name: str, units):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", "。", " ", ""],
    )
    chunks = []  # [{"text", "source", "location"}]
    for location, text in units:
        for piece in splitter.split_text(text):
            piece = piece.strip()
            if len(piece) > 20:
                chunks.append({"text": piece, "source": doc_name, "location": location})
    return chunks


@st.cache_resource(show_spinner=False)
def load_hf_model():
    """허깅페이스 로컬 임베딩 모델 로드 (한국어 특화 → 경량 다국어 순으로 시도)"""
    from sentence_transformers import SentenceTransformer

    last_err = None
    for name in HF_FALLBACK_MODELS:
        try:
            return name, SentenceTransformer(name)
        except Exception as e:
            last_err = e
    raise RuntimeError(f"허깅페이스 임베딩 모델 로드 실패: {last_err}")


def _embed_google(texts, task_type):
    if "api_key" in st.session_state and st.session_state["api_key"]:
        genai.configure(api_key=st.session_state["api_key"])
        
    # 1. 서버에 직접 물어봐서 현재 사용 가능한 임베딩 모델을 자동으로 찾습니다! (NotFound 원천 차단)
    available_models = [
        m.name for m in genai.list_models() 
        if 'embedContent' in m.supported_generation_methods
    ]
    
    if not available_models:
        raise RuntimeError("현재 API Key로 접근 가능한 구글 임베딩 모델이 없습니다. (API 권한 확인 필요)")
        
    # 2. 확인된 목록 중 가장 첫 번째(최신) 유효한 모델을 자동 지정합니다.
    target_model = available_models[0]
    
    embeddings = []
    BATCH = 90
    for i in range(0, len(texts), BATCH):
        batch = texts[i : i + BATCH]
        result = genai.embed_content(
            model=target_model,  # 👈 서버가 직접 확인한 확실한 모델명만 사용!
            content=batch,
            task_type=task_type
        )
        emb = result["embedding"]
        if emb and isinstance(emb[0], (int, float)):
            emb = [emb]
        embeddings.extend(emb)
    return embeddings


def _embed_hf(texts):
    name, model = load_hf_model()
    st.session_state.embed_model_name = name
    return model.encode(list(texts), show_progress_bar=False, normalize_embeddings=True).tolist()


def embed_texts(texts, task_type="retrieval_document"):
    """임베딩 생성.
    1순위: Google models/embedding-001
    Fallback: 허깅페이스 로컬 임베딩(jhgan/ko-sroberta-multitask → all-MiniLM-L6-v2)
    최초 성공한 백엔드를 세션에 고정하여 벡터 차원 혼용을 방지한다.
    """
    if isinstance(texts, str):
        texts = [texts]

    backend = st.session_state.get("embed_backend")

    if backend == "hf":
        return _embed_hf(texts)

    try:
        result = _embed_google(texts, task_type)
        if st.session_state.get("embed_backend") is None:
            st.session_state.embed_backend = "google"
            st.session_state.embed_model_name = EMBEDDING_MODEL
        return result
    except Exception as e:
        if backend == "google":
            # 이미 Google 벡터로 저장된 문서가 있으a74 차원 불일치 방지를 위해 전환 불가
            raise RuntimeError(
                f"Google 임베딩 API 호출 실패: {e}\n"
                "기존 문서가 Google 임베딩으로 저장되어 있어 벡터 차원 혼용을 막기 위해 중단합니다. "
                "'전체 문서 초기화' 후 다시 업로드하면 로컬 임베딩으로 자동 전환됩니다."
            )
        # 아직 아무 문서도 없으면 허깅페이스 로컬 임베딩으로 자동 전환
        st.session_state.embed_backend = "hf"
        st.toast(f"⚠️ Google 임베딩 실패({type(e).__name__}) → 로컬 허깅페이스 임베딩으로 자동 전환합니다.")
        return _embed_hf(texts)


def index_document(doc_hash: str, doc_name: str, file_bytes: bytes, progress_cb=None):
    units = parse_file(doc_name, file_bytes)
    if not units:
        raise ValueError("문서에서 텍스트를 추출하지 못했습니다.")
    chunks = chunk_units(doc_name, units)
    if not chunks:
        raise ValueError("유효한 텍스트 청크가 생성되지 않았습니다.")

    collection = get_collection()
    texts = [c["text"] for c in chunks]
    embeddings = embed_texts(texts, task_type="retrieval_document")

    ids = [f"{doc_hash}_{i}" for i in range(len(chunks))]
    metadatas = [{"source": c["source"], "location": c["location"]} for c in chunks]

    BATCH = 500
    for i in range(0, len(ids), BATCH):
        collection.add(
            ids=ids[i : i + BATCH],
            embeddings=embeddings[i : i + BATCH],
            documents=texts[i : i + BATCH],
            metadatas=metadatas[i : i + BATCH],
        )
    return len(chunks), len(units)


# ──────────────────────────────────────────────────────────────
# RAG 검색 + 답변 생성
# ──────────────────────────────────────────────────────────────
def retrieve(query: str, top_k: int = TOP_K):
    collection = get_collection()
    if collection.count() == 0:
        return []
    q_emb = embed_texts([query], task_type="retrieval_query")[0]
    n = min(top_k, collection.count())
    res = collection.query(query_embeddings=[q_emb], n_results=n)
    hits = []
    for doc, meta in zip(res["documents"][0], res["metadatas"][0]):
        hits.append({"text": doc, "source": meta["source"], "location": meta["location"]})
    return hits


def build_context(hits):
    blocks = []
    for i, h in enumerate(hits, 1):
        blocks.append(
            f"--- [참고 문서 {i}] (문서명: {h['source']}, 위치: {h['location']}) ---\n{h['text']}"
        )
    return "\n\n".join(blocks)


def generate_answer(model_name: str, question: str, hits):
    context = build_context(hits) if hits else "(검색된 참고 문서가 없습니다)"
    user_prompt = f"""[참고 문서]
{context}

[사용자 질문]
{question}

위의 [참고 문서]만을 근거로 규칙에 따라 답변하십시오."""

    # 지정 모델이 404(NotFound) 등으로 실패하면 폴백 후보를 순차 재시도
    tried, last_err = [], None
    for name in [model_name] + [m for m in GENERATION_FALLBACKS if m != model_name]:
        try:
            model = genai.GenerativeModel(
                model_name=name,
                system_instruction=SYSTEM_PROMPT,
                generation_config={"temperature": 0.0, "max_output_tokens": 8192},
            )
            response = model.generate_content(user_prompt)
            if name != model_name:
                st.caption(f"ℹ️ 선택 모델 호출 실패로 `{name}` 모델로 자동 대체 호출했습니다.")
            return response.text
        except Exception as e:
            tried.append(name)
            last_err = e
            # 모델 가용성 문제(404/NotFound/deprecated)일 때만 다음 후보 시도
            msg = str(e).lower()
            if not any(k in msg for k in ["404", "not found", "no longer available", "deprecated", "not supported"]):
                raise
    raise RuntimeError(f"사용 가능한 생성 모델이 없습니다 (시도: {tried}) — 마지막 오류: {last_err}")


# UI 라디오 선택값 → Google API가 인식하는 정확한 정식 모델명 매핑 (1:1 고정)
MODEL_MAP = {
    "gemini-1.5-pro (고정밀)": "gemini-1.5-pro",
    "gemini-1.5-flash (고속)": "gemini-1.5-flash",
}
# 선택 모델이 해당 계정에서 지원 종료/미제공인 경우에만 사용하는 비상용 폴백 후보
GENERATION_FALLBACKS = ["gemini-1.5-pro", "gemini-1.5-flash", "gemini-flash-latest", "gemini-2.0-flash"]


def resolve_model(selection: str):
    """UI 선택값 → 실제 호출 모델명 (정확한 정식 문자열 반환)"""
    primary = MODEL_MAP[selection]
    available = st.session_state.get("_available_models", set())
    if not available or f"models/{primary}" in available:
        return primary
    # 선택 모델이 계정에서 사용 불가한 경우, 사용 가능 목록에 있는 폴백 후보 반환
    for c in GENERATION_FALLBACKS:
        if f"models/{c}" in available:
            return c
    return primary


def configure_api(api_key: str) -> bool:
    try:
        genai.configure(api_key=api_key)
        models = {
            m.name
            for m in genai.list_models()
            if "generateContent" in getattr(m, "supported_generation_methods", [])
        }
        st.session_state["_available_models"] = models
        return True
    except Exception:
        return False


# ──────────────────────────────────────────────────────────────
# 사이드바 UI
# ──────────────────────────────────────────────────────────────
with st.sidebar:
    st.title("📚 연구 조교 설정")

    # 1) API Key
    st.subheader("1. Google Gemini API Key")
    env_key = os.getenv("GOOGLE_API_KEY", "") or os.getenv("GEMINI_API_KEY", "")
    api_key = st.text_input(
        "API Key 입력",
        value=env_key,
        type="password",
        help="환경변수(GOOGLE_API_KEY)가 없어도 여기에 직접 입력하면 됩니다. https://aistudio.google.com/apikey 에서 무료 발급 가능",
    )
    if api_key:
        os.environ["GOOGLE_API_KEY"] = api_key
        if not st.session_state.api_key_ok or st.session_state.get("_last_key") != api_key:
            with st.spinner("API Key 확인 중..."):
                ok = configure_api(api_key)
            st.session_state.api_key_ok = ok
            st.session_state["_last_key"] = api_key
        if st.session_state.api_key_ok:
            st.success("✅ API Key 인증 완료")
        else:
            st.error("❌ API Key가 유효하지 않습니다.")
    else:
        st.info("Gemini API Key를 입력해 주세요.")

    # 2) 모델 선택
    st.subheader("2. 답변 생성 모델")
    model_selection = st.radio(
        "모델 선택",
        ["gemini-1.5-pro (고정밀)", "gemini-1.5-flash (고속)"],
        index=0,
        help="선택값은 정식 모델명(gemini-1.5-pro / gemini-1.5-flash)으로 그대로 API에 전달됩니다. 해당 계정에서 모델이 제공되지 않는 경우에만 사용 가능한 모델로 자동 대체됩니다.",
    )

    st.divider()

    # 3) 파일 업로더
    st.subheader("3. 문서 업로드")
    uploaded_files = st.file_uploader(
        "TXT / PDF / EPUB / DOCX (다중 선택 가능)",
        type=["txt", "pdf", "epub", "docx"],
        accept_multiple_files=True,
    )

    if uploaded_files:
        if not st.session_state.api_key_ok and st.session_state.embed_backend != "hf":
            st.warning(
                "⚠️ 유효한 API Key가 없습니다. Google 임베딩 대신 로컬 허깅페이스 임베딩(Fallback)으로 진행할 수 있습니다. "
                "(단, 답변 생성에는 여전히 유효한 Gemini API Key가 필요합니다.)"
            )
            if st.button("💻 로컬 임베딩으로 계속하기", use_container_width=True):
                st.session_state.embed_backend = "hf"
                st.rerun()
        else:
            for uf in uploaded_files:
                fbytes = uf.getvalue()
                doc_hash = hashlib.md5(fbytes).hexdigest()[:12]
                if doc_hash in st.session_state.docs:
                    continue
                st.session_state.docs[doc_hash] = {
                    "name": uf.name, "status": "⏳ 처리 중", "chunks": 0, "pages": 0,
                }
                with st.spinner(f"'{uf.name}' 텍스트 추출 및 벡터화(임베딩) 중..."):
                    try:
                        n_chunks, n_units = index_document(doc_hash, uf.name, fbytes)
                        st.session_state.docs[doc_hash].update(
                            {"status": "✅ 완료", "chunks": n_chunks, "pages": n_units}
                        )
                    except Exception as e:
                        st.session_state.docs[doc_hash]["status"] = f"❌ 실패: {e}"

    # 4) 문서 상태표시줄
    st.subheader("4. 문서 벡터화 상태")
    if st.session_state.docs:
        done = sum(1 for d in st.session_state.docs.values() if d["status"].startswith("✅"))
        total = len(st.session_state.docs)
        st.progress(done / total if total else 0.0, text=f"벡터화 완료: {done} / {total}")
        for d in st.session_state.docs.values():
            st.markdown(
                f"- **{d['name']}**  \n"
                f"  상태: {d['status']}"
                + (f" · 청크 {d['chunks']}개 · 단위 {d['pages']}개" if d["chunks"] else "")
            )
        if st.button("🗑️ 전체 문서 초기화", use_container_width=True):
            try:
                if st.session_state.chroma_client is not None:
                    st.session_state.chroma_client.delete_collection(COLLECTION_NAME)
            except Exception:
                pass
            st.session_state.collection = None
            st.session_state.docs = {}
            st.session_state.messages = []
            st.session_state.embed_backend = None
            st.session_state.embed_model_name = None
            st.rerun()
    else:
        st.caption("아직 업로드된 문서가 없습니다.")

    if st.session_state.embed_backend:
        backend_label = (
            f"☁️ Google `{st.session_state.embed_model_name}`"
            if st.session_state.embed_backend == "google"
            else f"💻 로컬 HF `{st.session_state.embed_model_name}` (Fallback)"
        )
        st.caption(f"현재 임베딩 백엔드: {backend_label}")


# ──────────────────────────────────────────────────────────────
# 메인 화면
# ──────────────────────────────────────────────────────────────
st.title("📚 전천후 도서/논문 검증 및 연구 조교")
st.caption(
    "Google Gemini + RAG 기반 · 업로드한 문서에 명시된 사실만으로 답변하는 "
    "**환각 제로(Zero-Hallucination) 학술 검증 AI**입니다."
)

docs_ready = any(d["status"].startswith("✅") for d in st.session_state.docs.values())

# Quick Action Buttons
st.markdown("##### ⚡ 연구용 퀵 액션")
c1, c2, c3 = st.columns(3)
disabled = not (docs_ready and st.session_state.api_key_ok)
if c1.button("📌 서지 정보 추출", use_container_width=True, disabled=disabled):
    st.session_state.pending_prompt = QUICK_ACTIONS["bib"][1]
if c2.button("📊 핵심 요약", use_container_width=True, disabled=disabled):
    st.session_state.pending_prompt = QUICK_ACTIONS["summary"][1]
if c3.button("🔍 논리 교차 검증", use_container_width=True, disabled=disabled):
    st.session_state.pending_prompt = QUICK_ACTIONS["crosscheck"][1]

st.divider()

# 채팅 이력 렌더링
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])
        if msg.get("sources"):
            with st.expander("🔎 검색된 근거 문단 (RAG Retrieval)"):
                for s in msg["sources"]:
                    st.markdown(
                        f"> **[{s['source']} · {s['location']}]**  \n> {s['text'][:400]}..."
                    )

# 입력 처리
chat_input = st.chat_input("업로드한 문서에 대해 질문하세요. (예: 이 논문의 실험 방법은?)")
prompt = st.session_state.pending_prompt or chat_input
st.session_state.pending_prompt = None

if prompt:
    if not st.session_state.api_key_ok:
        st.error("먼저 좌측 사이드바에서 유효한 Gemini API Key를 입력해 주세요.")
    elif not docs_ready:
        st.error("먼저 좌측 사이드바에서 문서를 업로드하고 벡터화를 완료해 주세요.")
    else:
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)

        with st.chat_message("assistant"):
            try:
                with st.spinner("문서에서 근거를 검색하고 교차 검증 중..."):
                    hits = retrieve(prompt)
                    model_name = resolve_model(model_selection)
                    answer = generate_answer(model_name, prompt, hits)
                st.markdown(answer)

                # 출처 요약 표기
                if hits:
                    src_set = []
                    for h in hits:
                        tag = f"{h['source']} ({h['location']})"
                        if tag not in src_set:
                            src_set.append(tag)
                    st.caption("📎 참조된 근거 위치: " + " · ".join(src_set[:6]))
                    with st.expander("🔎 검색된 근거 문단 (RAG Retrieval)"):
                        for s in hits:
                            st.markdown(
                                f"> **[{s['source']} · {s['location']}]**  \n> {s['text'][:400]}..."
                            )
                st.session_state.messages.append(
                    {"role": "assistant", "content": answer, "sources": hits}
                )
            except Exception:
                err = traceback.format_exc(limit=2)
                st.error(f"답변 생성 중 오류가 발생했습니다:\n\n```\n{err}\n```")
